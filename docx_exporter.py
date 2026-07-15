from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

SUPPORTED_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg"}

COMMON_ITEM_FIELDS = [
    ("item_number", "Item Number"),
    ("evidence_number", "Evidence Number"),
    ("device_or_media_type", "Device / Media Type"),
    ("short_description", "Short Description"),
    ("condition_received", "Condition Received"),
    ("packaging_seal_info", "Packaging / Seal Information"),
    ("item_notes", "Item Notes"),
]

TYPE_SPECIFIC_LABELS = {
    "make_model": "Make / Model",
    "serial_number": "Serial Number",
    "imei_meid": "IMEI / MEID",
    "phone_number": "Phone Number",
    "carrier": "Carrier",
    "sim_present": "SIM Present",
    "storage_capacity": "Storage Capacity",
    "power_lock_state": "Power / Lock State",
    "passcode_provided": "Passcode / Password Provided",
    "known_account_info": "Known Account Info",
    "serial_service_tag": "Serial / Service Tag",
    "operating_system": "Operating System",
    "storage_type": "Storage Type",
    "power_state": "Power State",
    "login_credentials_provided": "Login Credentials Provided",
    "known_user_account": "Known User Account",
    "brand_model": "Brand / Model",
    "connector_type": "Connector Type",
    "encryption_suspected": "Encryption Suspected",
    "card_type": "Card Type",
    "adapter_included": "Adapter Included",
    "connection_type": "Connection Type",
    "power_supply_included": "Power Supply Included",
    "channel_count": "Channel Count",
    "date_time_setting": "Date / Time Setting",
    "export_format": "Export Format",
    "network_info": "Network Info",
    "platform_provider": "Platform / Provider",
    "account_identifier": "Account Identifier",
    "return_export_date": "Return / Export Date",
    "source_authority": "Source Authority",
    "file_folder_location": "File / Folder Location",
    "export_date": "Export Date",
    "date_range": "Date Range",
    "producing_party": "Producing Party",
    "media_type": "Media Type",
    "source_device_platform": "Source Device / Platform",
    "file_folder_path": "File / Folder Path",
    "provided_by": "Provided By",
    "original_source_known": "Original Source Known",
    "identifier": "Identifier",
    "description": "Description",
    "capacity_size": "Capacity / Size",
    "source_origin": "Source / Origin",
    "additional_details": "Additional Details",
}

LEGACY_TYPE_FIELD_MAP = {
    "make_model": "make_model",
    "serial_number": "serial_number",
    "imei_meid": "imei_meid",
    "phone_number": "phone_number",
    "storage_capacity": "storage_capacity",
    "power_lock_state": "power_lock_state",
    "passcode_provided": "passcode_provided",
    "known_account_info": "known_account_info",
}


def set_document_defaults(document):
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.bold = True


def add_branding_image(document, settings):
    path_value = settings.get("report_branding", {}).get("patch_image_path", "").strip()
    if not path_value:
        return

    image_path = Path(path_value)
    if not image_path.exists() or not image_path.is_file():
        return

    if image_path.suffix.lower() not in SUPPORTED_IMAGE_EXTENSIONS:
        return

    try:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(image_path), width=Inches(1.1))
    except Exception:
        return


def add_key_value_table(document, rows):
    filtered_rows = [(label, value) for label, value in rows if str(value or "").strip()]

    if not filtered_rows:
        document.add_paragraph("No information provided.")
        return

    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    for label, value in filtered_rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value or "")

    document.add_paragraph()


def add_key_value_table_allow_empty(document, rows):
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value or "")

    document.add_paragraph()


def add_bullets(document, values):
    if values:
        for value in values:
            document.add_paragraph(str(value), style="List Bullet")
    else:
        document.add_paragraph("None selected.")


def add_text_section(document, heading, text):
    document.add_heading(heading, level=1)
    document.add_paragraph(text or "")


def get_type_specific(item):
    type_specific = item.get("type_specific", {})

    if not isinstance(type_specific, dict):
        type_specific = {}

    cleaned = {}

    for key, value in type_specific.items():
        value = str(value or "").strip()
        if value:
            cleaned[str(key).strip()] = value

    for legacy_key, type_key in LEGACY_TYPE_FIELD_MAP.items():
        value = str(item.get(legacy_key, "")).strip()
        if value and type_key not in cleaned:
            cleaned[type_key] = value

    return cleaned


def get_item_rows(item):
    rows = []

    for key, label in COMMON_ITEM_FIELDS:
        value = item.get(key, "")
        if value:
            rows.append((label, value))

    type_specific = get_type_specific(item)

    for key, value in type_specific.items():
        rows.append((TYPE_SPECIFIC_LABELS.get(key, key.replace("_", " ").title()), value))

    return rows


def add_item_photos(document, photos):
    document.add_heading("Item Photos", level=3)

    if not photos:
        document.add_paragraph("No item photos listed.")
        return

    for photo in photos:
        document.add_paragraph(f"Photo {photo.get('photo_number', '')}", style="List Bullet")
        add_key_value_table(document, [
            ("Type", photo.get("photo_type", "")),
            ("File Name", photo.get("file_name", "")),
            ("Relative Copied Path", photo.get("relative_copied_path", "")),
            ("Copy Status", photo.get("copy_status", "")),
            ("Copy Error", photo.get("copy_error", "")),
            ("Description", photo.get("description", "")),
        ])


def add_peripherals(document, peripherals):
    document.add_heading("Peripherals / Accessories", level=3)

    if not peripherals:
        document.add_paragraph("No peripherals/accessories listed.")
        return

    for peripheral in peripherals:
        heading = f"Peripheral {peripheral.get('peripheral_number', '')}"
        if peripheral.get("peripheral_type"):
            heading += f" - {peripheral.get('peripheral_type')}"
        document.add_paragraph(heading, style="List Bullet")
        add_key_value_table(document, [
            ("Type", peripheral.get("peripheral_type", "")),
            ("Description", peripheral.get("description", "")),
            ("Serial / Identifier", peripheral.get("serial_identifier", "")),
            ("Condition", peripheral.get("condition", "")),
            ("Included With Item", peripheral.get("included_with_item", "")),
            ("Notes", peripheral.get("notes", "")),
        ])


def add_evidence_items(document, items):
    document.add_heading("Evidence / Device Items", level=1)

    if not items:
        document.add_paragraph("No evidence/device items listed.")
        return

    for index, item in enumerate(items, start=1):
        item_number = item.get("item_number", "") or str(index).zfill(3)
        device_type = item.get("device_or_media_type", "")
        heading_text = f"Item {item_number}"

        if device_type:
            heading_text += f" - {device_type}"

        document.add_heading(heading_text, level=2)
        add_key_value_table(document, get_item_rows(item))
        add_peripherals(document, item.get("peripherals", []))
        add_item_photos(document, item.get("item_photos", []))


def add_attachments(document, attachments):
    document.add_heading("Attachment Index", level=1)

    if not attachments:
        document.add_paragraph("No supporting documents listed.")
        return

    for attachment in attachments:
        number = attachment.get("attachment_number", "")
        attachment_type = attachment.get("attachment_type", "")
        heading_text = f"Attachment {number}"

        if attachment_type:
            heading_text += f" - {attachment_type}"

        document.add_heading(heading_text, level=2)

        add_key_value_table(document, [
            ("Type", attachment.get("attachment_type", "")),
            ("File Name", attachment.get("file_name", "")),
            ("Relative Copied Path", attachment.get("relative_copied_path", "")),
            ("Copy Status", attachment.get("copy_status", "")),
            ("Copy Error", attachment.get("copy_error", "")),
            ("Related Item", attachment.get("related_item", "")),
            ("Document Date", attachment.get("document_date", "")),
            ("Provided By", attachment.get("provided_by", "")),
            ("Description", attachment.get("description", "")),
            ("Notes", attachment.get("notes", "")),
        ])


def add_handoff(document, handoff):
    document.add_heading("Submission / Handoff", level=1)
    add_key_value_table(document, [
        ("Submitted By", handoff.get("submitted_by", "")),
        ("Submitted By Title / Unit", handoff.get("submitted_by_title_unit", "")),
        ("Submitted Date / Time", handoff.get("submitted_date_time", "")),
        ("Submitted To", handoff.get("submitted_to", "")),
        ("Receiving Unit / Lab", handoff.get("receiving_unit_lab", "")),
        ("Transfer Method", handoff.get("transfer_method", "")),
        ("Evidence Packaging / Seal Number", handoff.get("packaging_seal_info", "")),
        ("Condition at Submission", handoff.get("condition_at_submission", "")),
        ("Evidence Items Submitted", handoff.get("item_count_submitted", "")),
        ("Attachments Listed", handoff.get("attachment_count_submitted", "")),
        ("Attachments Copied", handoff.get("attachments_copied", "")),
        ("Item Photos Listed", handoff.get("item_photo_count", "")),
        ("Item Photos Copied", handoff.get("item_photos_copied", "")),
        ("Peripherals / Accessories Listed", handoff.get("peripheral_count", "")),
        ("Received By", handoff.get("received_by", "")),
        ("Received Date / Time", handoff.get("received_date_time", "")),
        ("Receiving Notes", handoff.get("receiving_notes", "")),
    ])


def add_acknowledgement(document, handoff, investigator):
    document.add_heading("Submission Acknowledgement", level=1)
    document.add_paragraph(
        "This acknowledgement documents receipt of a digital forensic request packet and related submitted materials. "
        "It does not replace the agency's official evidence management system, chain-of-custody record, property record, "
        "legal review process, official evidence photographs, or forensic examination report."
    )
    add_key_value_table_allow_empty(document, [
        ("Submitted By", handoff.get("submitted_by") or investigator.get("submitting_investigator", "")),
        ("Title / Unit", handoff.get("submitted_by_title_unit", "")),
        ("Date / Time Submitted", handoff.get("submitted_date_time", "")),
        ("Received By", handoff.get("received_by", "")),
        ("Receiving Unit / Lab", handoff.get("receiving_unit_lab", "")),
        ("Date / Time Received", handoff.get("received_date_time", "")),
        ("Items Submitted", handoff.get("item_count_submitted", "")),
        ("Attachments Listed", handoff.get("attachment_count_submitted", "")),
        ("Item Photos Listed", handoff.get("item_photo_count", "")),
        ("Packaging / Seal Information", handoff.get("packaging_seal_info", "")),
    ])
    document.add_paragraph("Submitted By Signature: _______________________________")
    document.add_paragraph("Receiving Staff Signature: _______________________________")


def build_docx_request(packet, settings):
    document = Document()
    set_document_defaults(document)
    add_branding_image(document, settings)

    agency = packet.get("agency", {})
    case_info = packet.get("case_info", {})
    investigator = packet.get("investigator_info", {})
    authority = packet.get("legal_authority", {})
    scope = packet.get("scope", {})
    items = packet.get("evidence_items", [])
    attachments = packet.get("attachments", [])
    details = packet.get("request_details", {})
    priority = packet.get("priority_info", {})
    handoff = packet.get("handoff_info", {})
    options = packet.get("report_options", {})

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BYTECASE INTAKE")
    run.bold = True
    run.font.size = Pt(16)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Digital Forensics Request Packet")

    if agency.get("agency_name") or agency.get("unit_name"):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if agency.get("agency_name"):
            paragraph.add_run(agency.get("agency_name", ""))
        if agency.get("agency_name") and agency.get("unit_name"):
            paragraph.add_run("\n")
        if agency.get("unit_name"):
            paragraph.add_run(agency.get("unit_name", ""))

    document.add_paragraph()

    document.add_heading("Request Information", level=1)
    add_key_value_table_allow_empty(document, [
        ("Generated By", f"{packet.get('app_name', '')} - {packet.get('app_subtitle', '')} v{packet.get('app_version', '')}"),
        ("Created At", packet.get("created_at", "")),
        ("Schema", f"{packet.get('schema_name', '')} v{packet.get('schema_version', '')}"),
        ("Product Domain", packet.get("product_domain", "")),
    ])

    document.add_heading("Case Information", level=1)
    add_key_value_table(document, [
        ("Case Number", case_info.get("case_number", "")),
        ("Agency Case Number", case_info.get("agency_case_number", "")),
        ("Offense / Incident Type", case_info.get("offense_type", "")),
    ])

    document.add_heading("Submitting Investigator", level=1)
    add_key_value_table(document, [
        ("Submitting Agency", investigator.get("submitting_agency", "")),
        ("Submitting Unit", investigator.get("submitting_unit", "")),
        ("Submitting Investigator", investigator.get("submitting_investigator", "")),
        ("Phone", investigator.get("investigator_phone", "")),
        ("Email", investigator.get("investigator_email", "")),
        ("Supervisor", investigator.get("supervisor", "")),
    ])

    document.add_heading("Legal Authority", level=1)
    add_key_value_table(document, [
        ("Authority Type", authority.get("authority_type", "")),
        ("Authority Date", authority.get("authority_date", "")),
        ("Authority Reference / Document Number", authority.get("authority_reference", "")),
        ("Authority Notes", authority.get("authority_notes", "")),
    ])

    document.add_heading("Legal Authority / Scope Options", level=1)
    add_bullets(document, scope.get("scope_options", []))

    document.add_heading("Scope Limitations", level=1)
    add_key_value_table(document, [
        ("Date Range Limited", "Yes" if scope.get("date_range_limited") else "No"),
        ("Authorized Date Range", scope.get("authorized_date_range", "")),
        ("Specific Apps / Platforms Limited", "Yes" if scope.get("specific_apps_limited") else "No"),
        ("Authorized Apps / Platforms", scope.get("authorized_apps_platforms", "")),
        ("Specific Accounts Limited", "Yes" if scope.get("specific_accounts_limited") else "No"),
        ("Authorized Accounts", scope.get("authorized_accounts", "")),
        ("Keyword Limited", "Yes" if scope.get("keyword_limited") else "No"),
        ("Authorized Keywords", scope.get("authorized_keywords", "")),
        ("Additional Scope Notes", scope.get("scope_notes", "")),
    ])

    add_evidence_items(document, items)
    add_attachments(document, attachments)
    document.add_heading("Requested Forensic Actions", level=1)
    add_bullets(document, details.get("requested_actions", []))

    add_text_section(document, "Investigative Objective / Requested Information", details.get("investigative_objective", ""))
    add_text_section(document, "Known Facts / Context for Examiner", details.get("known_facts_for_examiner", ""))

    document.add_heading("Priority / Urgency", level=1)
    add_key_value_table(document, [
        ("Priority", priority.get("priority", "")),
        ("Requested Due Date", priority.get("requested_due_date", "")),
    ])
    document.add_paragraph("Urgency Flags:")
    add_bullets(document, priority.get("urgency_flags", []))
    add_text_section(document, "Priority Notes", priority.get("priority_notes", ""))

    add_handoff(document, handoff)

    document.add_heading("Important Scope Note", level=1)
    document.add_paragraph(packet.get("disclaimer", ""))

    if options.get("include_acknowledgement_block"):
        add_acknowledgement(document, handoff, investigator)

    document.add_paragraph()
    end = document.add_paragraph("End of ByteCase Intake Request Packet")
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return document


def save_docx_request(packet, settings, docx_path):
    document = build_docx_request(packet, settings)
    document.save(docx_path)
    return docx_path
